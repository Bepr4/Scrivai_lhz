"""DocxRenderer — docxtpl 模板渲染。

约束(docxtpl 限制):
1. 模板必须由 Word/LibreOffice 手工制作(jinja 标签必须在单一 <w:r>)
2. 单 cell 内不支持嵌套 {% for %};用 jinja 过滤器扁平化
3. 避免表中表
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docxtpl import DocxTemplate

# 正则匹配 docxtpl 占位符 {{ var }} 中的 var 名
_PLACEHOLDER_RE = re.compile(r"\{\{\s*([a-zA-Z_][a-zA-Z0-9_]*)\s*\}\}")


class DocxRenderer:
    """基于 docxtpl 的 docx 模板渲染器。"""

    def __init__(self, template_path: str | Path) -> None:
        self._template_path = Path(template_path)
        if not self._template_path.is_file():
            raise FileNotFoundError(f"docx 模板不存在:{self._template_path}")
        # 加载一次校验文件能被 docxtpl 解析
        self._template = DocxTemplate(str(self._template_path))

    @property
    def template_path(self) -> Path:
        return self._template_path

    def list_placeholders(self) -> list[str]:
        """正则扫描模板内全部 {{ var }} 占位符;返回去重排序的 var 名列表。"""
        # docxtpl 内部 docx XML 的纯文本提取
        names: set[str] = set()
        # 重新 open 拿 raw text(不依赖 docxtpl 的 jinja env)
        from docx import Document

        doc = Document(str(self._template_path))
        for para in doc.paragraphs:
            names.update(_PLACEHOLDER_RE.findall(para.text))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        names.update(_PLACEHOLDER_RE.findall(para.text))
        return sorted(names)

    def render(self, context: dict[str, Any], output_path: str | Path) -> Path:
        """渲染模板并写到 output_path;失败不留半成品。

        每次 render 重新加载模板(DocxTemplate.render 是一次性消费的)。
        """
        out = Path(output_path)
        # 确保父目录存在;若父目录都不存在则报错(测试期望)
        if not out.parent.is_dir():
            raise IOError(f"输出目录不存在:{out.parent}")

        # 重新 open 模板(DocxTemplate 渲染是 stateful 的)
        tpl = DocxTemplate(str(self._template_path))
        try:
            tpl.render(context)
            tpl.save(str(out))
        except Exception:
            # 任何异常:删半成品再 re-raise
            if out.exists():
                try:
                    out.unlink()
                except OSError:
                    pass
            raise
        return out
