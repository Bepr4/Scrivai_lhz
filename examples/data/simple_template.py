"""构造 docxtpl 最小模板到 OUT_PATH。

单独抽出来是为了避免在 examples/data/ 入库二进制 .docx。
"""

from __future__ import annotations

from pathlib import Path

from docx import Document


def build_template(out_path: Path) -> Path:
    """构造含 3 个占位符的 docxtpl 模板。

    参数:
        out_path: 输出 .docx 路径,父目录会自动创建。

    返回:
        out_path。
    """
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("{{project_name}} 工程概况", level=1)
    doc.add_paragraph("工程地点:{{project_location}}")
    doc.add_paragraph("工程规模:{{project_scale}}")
    doc.save(str(out_path))
    return out_path


if __name__ == "__main__":
    import sys

    target = Path(sys.argv[1] if len(sys.argv) > 1 else "/tmp/scrivai-examples/simple_template.docx")
    path = build_template(target)
    print(f"[OK] 模板已生成: {path}")
